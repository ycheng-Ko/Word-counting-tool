import sys
import docx
import datetime

"""Define Functions"""
def readfile(fname):
    f = open(fname, "r")
    f.lstrip

    return 

# function to seperate paragraphs
def para_cutter(f): 
    para = []
    
    return para
"""Main"""
fname = sys.argv[1]
#outtype = sys.argv[2]

para_list = readfile( fname )
para_num = 0
total_num = 0
today = str(datetime.datetime.now().month) + str(datetime.datetime.now().day)
output = docx.Document()

for i in range( len(para_list) ):   
    word =  para_list[i].split()
    if len(word) > 1:
        if i < len(para_list) - 1:          # Other paragraph
            output.add_paragraph( para_list[i]) #, end='')
            para_num += len(word)
        else:                               # The last paragraph
            output.add_paragraph( para_list[i])#, end='')
            para_num += len(word)
            output.add_paragraph(f"({para_num})")
            total_num += para_num
    else:
        if i > 1:
            output.add_paragraph(f"({para_num})\n")
            total_num += para_num
            para_num = 0
        
output.add_paragraph(f"\n({total_num} words)")

if sys.argv[2] != None:
    if sys.argv[2] == "-i":
        output.save("iss" + today + ".docx")
    elif sys.argv[2] == "-a":
        output.save("arg" + today + ".docx")
else:
    output(today + ".docx")